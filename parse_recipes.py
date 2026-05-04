#!/usr/bin/env python3
"""
parse_recipes.py — Parse all .docx recipe files from Joanne's Cookbook
into a SQLite database with full-text search support.

Usage:
    python3 parse_recipes.py
"""

import os
import re
import sqlite3
import json
import sys

import docx
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
COOKBOOK_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Joanne_s Cookbook")
DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "recipes.db")

SKIP_FILES = {
    "open file for future use.docx",
    "open space for future additions.docx",
}

# Threshold in EMU – 16 pt = 203200 EMU.  Section headers are ≥ 16 pt bold
# but *not* underlined.
SECTION_HEADER_SIZE_THRESHOLD = 190500  # ~15 pt — safe lower bound

# ---------------------------------------------------------------------------
# Database helpers
# ---------------------------------------------------------------------------

def init_db(db_path: str) -> sqlite3.Connection:
    """Create (or recreate) the database and return a connection."""
    if os.path.exists(db_path):
        os.remove(db_path)
    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.executescript(
        """
        CREATE TABLE recipes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            category TEXT,
            subcategory TEXT,
            ingredients TEXT,
            instructions TEXT,
            notes TEXT,
            serves TEXT,
            source_file TEXT
        );

        CREATE VIRTUAL TABLE recipes_fts USING fts5(
            title, category, subcategory, ingredients, instructions, notes,
            content='recipes', content_rowid='id'
        );

        -- Triggers to keep FTS in sync
        CREATE TRIGGER recipes_ai AFTER INSERT ON recipes BEGIN
            INSERT INTO recipes_fts(rowid, title, category, subcategory,
                                    ingredients, instructions, notes)
            VALUES (new.id, new.title, new.category, new.subcategory,
                    new.ingredients, new.instructions, new.notes);
        END;

        CREATE TRIGGER recipes_ad AFTER DELETE ON recipes BEGIN
            INSERT INTO recipes_fts(recipes_fts, rowid, title, category,
                                    subcategory, ingredients, instructions, notes)
            VALUES ('delete', old.id, old.title, old.category, old.subcategory,
                    old.ingredients, old.instructions, old.notes);
        END;

        CREATE TRIGGER recipes_au AFTER UPDATE ON recipes BEGIN
            INSERT INTO recipes_fts(recipes_fts, rowid, title, category,
                                    subcategory, ingredients, instructions, notes)
            VALUES ('delete', old.id, old.title, old.category, old.subcategory,
                    old.ingredients, old.instructions, old.notes);
            INSERT INTO recipes_fts(rowid, title, category, subcategory,
                                    ingredients, instructions, notes)
            VALUES (new.id, new.title, new.category, new.subcategory,
                    new.ingredients, new.instructions, new.notes);
        END;
        """
    )
    return conn


def insert_recipe(conn, recipe: dict):
    conn.execute(
        """INSERT INTO recipes
           (title, category, subcategory, ingredients, instructions, notes, serves, source_file)
           VALUES (:title, :category, :subcategory, :ingredients, :instructions, :notes, :serves, :source_file)""",
        recipe,
    )


# ---------------------------------------------------------------------------
# Document element iteration (handles interleaved paragraphs + tables)
# ---------------------------------------------------------------------------

def iter_body_elements(doc):
    """Yield (type, element) tuples in document-body order."""
    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p" and child in para_map:
            yield ("p", para_map[child])
        elif tag == "tbl" and child in table_map:
            yield ("tbl", table_map[child])


# ---------------------------------------------------------------------------
# Paragraph classification helpers
# ---------------------------------------------------------------------------

def _is_large_font(run) -> bool:
    """Check if a run uses a large font (section header size)."""
    return bool(run.font.size and run.font.size >= SECTION_HEADER_SIZE_THRESHOLD)


def is_recipe_title(para) -> bool:
    """A recipe title is detected by ANY of these patterns:

    Pattern A: Text contains '___' (with non-empty cleaned text)
               — the most reliable marker, regardless of bold/underline
    Pattern B: Any run is both bold AND underlined (without '___')
               — catches titles where author forgot trailing underscores

    Excludes:
      - Empty / whitespace-only paragraphs
      - Large-font section headers (>=15pt bold, no underscores)
      - Very long text (>120 chars after stripping underscores) — likely not a title
    """
    text = para.text.strip()
    if not text:
        return False

    cleaned = text.rstrip("_").strip()
    if not cleaned:
        return False

    # Reject very long lines — recipe titles are short
    if len(cleaned) > 120:
        return False

    runs_with_text = [r for r in para.runs if r.text.strip().rstrip("_").strip()]
    if not runs_with_text:
        return False

    has_underscores = "___" in text
    has_any_underline = any(r.underline for r in para.runs if r.text.strip())
    has_any_bold = any(r.bold for r in para.runs if r.text.strip())
    has_bold_underline = any(
        r.bold and r.underline for r in para.runs if r.text.strip().rstrip("_").strip()
    )
    is_heading = para.style.name.lower().startswith("heading")

    # Reject large-font section headers (bold, big, no underscores)
    if not has_underscores:
        for run in runs_with_text:
            if _is_large_font(run):
                return False

    # Pattern A: Has '___' — accept regardless of bold/underline
    # This catches all formatting variants: bold+UL, UL-only, heading-style,
    # and even runs where underline was lost
    if has_underscores:
        # Still reject large-font section headers that happen to have ___
        for run in runs_with_text:
            if _is_large_font(run) and "___" not in run.text:
                return False
        return True

    # Pattern B: Bold + Underline without '___'
    # Author forgot underscores but formatted correctly
    if has_bold_underline:
        return True

    # Pattern C: Bold-only (no underline, no ___)
    # Only if ALL runs are bold AND text looks like a recipe title
    # (not a section label, not a sub-component header)
    if has_any_bold and all(r.bold for r in runs_with_text):
        # Reject section labels that end with ':' and are short
        if text.endswith(':') and len(cleaned) < 30:
            return False
        # Reject known sub-component / section patterns
        lower = cleaned.lower().rstrip(':')
        # Prefix patterns — reject if title starts with these
        skip_prefixes = [
            'for the', 'ideas for',
            'cont:', 'continued', 'step ', 'see ', 'also ',
            'base recipe',
        ]
        # Exact-match patterns — reject only if title IS this word
        skip_exact = {
            'crust', 'filling', 'topping', 'frosting', 'glaze', 'icing',
            'garnish', 'assembly', 'batter', 'marinade', 'coating', 'broth',
            'custard', 'whipped cream', 'additions', 'changes',
            'directions', 'method', 'preparation',
            'notes', 'variations', 'tips', 'serving',
            'optional', 'possible', 'other',
            'ingredients', 'instructions', 'cooking',
            'sauce', 'meat', 'cheese', 'lobster', 'ciders', 'salsa',
            'cakes', 'plum', 'fillings',
        }
        if lower in skip_exact:
            return False
        if any(lower.startswith(p) for p in skip_prefixes):
            return False
        # Reject very short generic words (but keep things like GHEE)
        if len(cleaned) < 4:
            return False
        # Reject numbered section headers like '1.CHEESE', '2. FILLINGS'
        if re.match(r'^\d+\.\s*', cleaned):
            return False
        # Reject category/TOC-style headers (parenthetical descriptions)
        if cleaned.startswith('(') and cleaned.endswith(')'):
            return False
        # Accept: this is likely a bold recipe title
        return True

    # Pattern D: Underline-only (no bold, no ___)
    # Some recipes have underlined titles without bold formatting
    if has_any_underline and not has_any_bold and len(cleaned) >= 4 and len(cleaned) <= 80:
        return True

    # Pattern E: Heading-style paragraphs
    # Word Heading styles inherit bold/underline from the style definition,
    # but run.bold returns None ("inherited") rather than True.
    # Check for Heading styles with short, title-like text.
    # Reject ingredient lines (start with - or •) that happen to have heading style
    if is_heading and len(cleaned) >= 4 and len(cleaned) <= 80 and not cleaned.startswith('-') and not cleaned.startswith('•'):
        # Reject generic section labels
        lower_h = cleaned.lower().rstrip(':')
        heading_skip = {
            'ingredients', 'instructions', 'directions', 'notes',
            'method', 'preparation', 'variations', 'tips',
        }
        if lower_h not in heading_skip:
            return True

    return False


def is_section_header(para) -> bool:
    """Section headers are bold, large font (>=16pt), but NOT underlined."""
    text = para.text.strip()
    if not text:
        return False
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if not runs_with_text:
        return False
    for run in runs_with_text:
        if not run.bold:
            return False
        if run.font.size and run.font.size >= SECTION_HEADER_SIZE_THRESHOLD:
            return True
    return False


def clean_title(text: str) -> str:
    """Strip trailing underscores and whitespace from a recipe title."""
    return text.strip().rstrip("_").strip()


def is_ingredient_line(text: str) -> bool:
    """Lines starting with -, •, or tab-indented dashes."""
    stripped = text.lstrip()
    if stripped.startswith("-") or stripped.startswith("•"):
        return True
    if stripped.startswith("\t-") or stripped.startswith("\t•"):
        return True
    return False


_INGREDIENT_LABEL_RE = re.compile(
    r"^Ingredients\s*[;:]", re.IGNORECASE
)

_SERVES_RE = re.compile(
    r"(?:serves|makes|yield[s]?)\s+(.+)", re.IGNORECASE
)

_NOTE_RE = re.compile(
    r"^(?:NOTES?|SUBSTITUT(?:ION|ITION)S?)\s*[;:.]\s*", re.IGNORECASE
)


def extract_serves(text: str) -> str | None:
    """Pull serving info from an Ingredients: line."""
    m = _SERVES_RE.search(text)
    if m:
        return m.group(1).strip().rstrip(".")
    return None


# ---------------------------------------------------------------------------
# Table text extraction
# ---------------------------------------------------------------------------

def table_to_lines(table) -> list[str]:
    """Convert a Word table into pseudo-ingredient lines."""
    lines = []
    for row in table.rows:
        parts = [cell.text.strip() for cell in row.cells]
        deduped = []
        for p in parts:
            if not deduped or p != deduped[-1]:
                deduped.append(p)
        line = " ".join(deduped).strip()
        if line:
            lines.append("- " + line)
    return lines


# ---------------------------------------------------------------------------
# Core parser
# ---------------------------------------------------------------------------

def parse_docx(filepath: str, category: str, subcategory: str) -> list[dict]:
    """Parse a single .docx file and return a list of recipe dicts."""
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"  WARNING: Could not open {filepath}: {e}")
        return []

    # Collect all content in body order (paragraphs + table text)
    elements: list[tuple[str, object]] = []
    for etype, elem in iter_body_elements(doc):
        if etype == "p":
            elements.append(("p", elem))
        elif etype == "tbl":
            for line in table_to_lines(elem):
                elements.append(("text", line))

    # -------------------------------------------------------------------
    # Pass 1: Identify title boundaries
    # -------------------------------------------------------------------
    title_indices: list[int] = []
    for i, (etype, elem) in enumerate(elements):
        if etype == "p" and is_recipe_title(elem):
            title_indices.append(i)

    if not title_indices:
        return []

    # -------------------------------------------------------------------
    # Pass 2: Extract recipe blocks between titles
    # -------------------------------------------------------------------
    recipes = []
    for idx, ti in enumerate(title_indices):
        etype, elem = elements[ti]
        title = clean_title(elem.text)
        if not title:
            continue

        # Collect content paragraphs from after title to next title (or end)
        end = title_indices[idx + 1] if idx + 1 < len(title_indices) else len(elements)
        content_lines: list[str] = []
        for j in range(ti + 1, end):
            ctype, celem = elements[j]
            if ctype == "p":
                content_lines.append(celem.text)
            else:  # "text" from table
                content_lines.append(celem)

        recipe = _parse_recipe_content(title, content_lines, category, subcategory, filepath)
        recipes.append(recipe)

    # Post-filter: remove single-word sub-component titles that slipped
    # through Pattern A/B (they have underscores or bold+underline but are
    # just sub-components of the previous recipe, e.g. "Glaze", "Frosting")
    _SUB_COMPONENT_EXACT = {
        'glaze', 'frosting', 'icing', 'filling', 'topping', 'crust',
        'sauce', 'garnish', 'batter', 'coating', 'marinade', 'custard',
        'assembly', 'broth', 'salsa', 'cake', 'pork',
    }
    recipes = [
        r for r in recipes
        if r['title'].lower().strip() not in _SUB_COMPONENT_EXACT
    ]

    return recipes


def _parse_recipe_content(
    title: str,
    lines: list[str],
    category: str,
    subcategory: str,
    filepath: str,
) -> dict:
    """Classify content lines into ingredients, instructions, and notes."""
    ingredients: list[str] = []
    instructions: list[str] = []
    notes: list[str] = []
    serves: str | None = None

    state = "pre"

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if state == "ing":
                state = "inst"
            continue

        # Check for Ingredients: label
        if _INGREDIENT_LABEL_RE.match(line):
            s = extract_serves(line)
            if s:
                serves = s
            state = "ing"
            continue

        # Check for notes
        if _NOTE_RE.match(line):
            notes.append(line)
            state = "notes"
            continue

        if line.upper().startswith("SUBSTITUT") or line.upper().startswith("SUBSTITUITI"):
            notes.append(line)
            state = "notes"
            continue

        if state == "notes":
            if is_ingredient_line(line) or _INGREDIENT_LABEL_RE.match(line):
                state = "ing"
            else:
                notes.append(line)
                continue

        if state in ("pre", "ing"):
            if is_ingredient_line(line):
                state = "ing"
                ingredients.append(line)
                continue
            elif state == "ing":
                state = "inst"
                instructions.append(line)
                continue
            else:
                instructions.append(line)
                continue

        if state == "inst":
            if is_ingredient_line(line):
                ingredients.append(line)
                continue
            instructions.append(line)

    return {
        "title": title,
        "category": category,
        "subcategory": subcategory,
        "ingredients": "\n".join(ingredients) if ingredients else None,
        "instructions": "\n".join(instructions) if instructions else None,
        "notes": "\n".join(notes) if notes else None,
        "serves": serves,
        "source_file": os.path.relpath(filepath, os.path.dirname(os.path.abspath(__file__))),
    }


# ---------------------------------------------------------------------------
# File discovery
# ---------------------------------------------------------------------------

def discover_docx_files(base_dir: str) -> list[tuple[str, str, str]]:
    """Return list of (filepath, category, subcategory) tuples."""
    results = []
    base_dir = os.path.abspath(base_dir)

    for root, _dirs, files in os.walk(base_dir):
        for fname in sorted(files):
            if not fname.lower().endswith(".docx"):
                continue
            if fname.lower() in {s.lower() for s in SKIP_FILES}:
                continue

            filepath = os.path.join(root, fname)
            rel = os.path.relpath(root, base_dir)

            if rel == ".":
                category = ""
            else:
                parts = rel.split(os.sep)
                category = parts[0].strip()

            subcategory = os.path.splitext(fname)[0].strip()
            results.append((filepath, category, subcategory))

    return results


# ---------------------------------------------------------------------------
# Export to JSON for static site
# ---------------------------------------------------------------------------

def export_json(db_path: str, json_path: str):
    """Export recipes from SQLite to compact JSON for the static frontend."""
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    rows = conn.execute(
        "SELECT id, title, category, subcategory, ingredients, instructions, notes, serves FROM recipes"
    ).fetchall()
    recipes = []
    for r in rows:
        d = dict(r)
        for k in d:
            if d[k] is None:
                d[k] = ""
        recipes.append(d)
    conn.close()

    with open(json_path, "w") as f:
        json.dump(recipes, f, separators=(",", ":"))

    print(f"Exported {len(recipes)} recipes to {json_path}")
    size_mb = os.path.getsize(json_path) / 1024 / 1024
    print(f"JSON file size: {size_mb:.1f} MB")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    base_dir = os.path.abspath(COOKBOOK_DIR)
    db_path = os.path.abspath(DB_PATH)
    json_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "public", "recipes.json")

    print(f"Cookbook directory: {base_dir}")
    print(f"Database path:     {db_path}")
    print()

    files = discover_docx_files(base_dir)
    print(f"Found {len(files)} .docx files to parse.")
    print()

    conn = init_db(db_path)
    total_recipes = 0

    for i, (filepath, category, subcategory) in enumerate(files, 1):
        short = os.path.relpath(filepath, base_dir)
        recipes = parse_docx(filepath, category, subcategory)
        for r in recipes:
            insert_recipe(conn, r)
        total_recipes += len(recipes)
        print(f"  [{i:3d}/{len(files)}] {short:<65s}  → {len(recipes):3d} recipes")

    conn.commit()

    # Verify
    count = conn.execute("SELECT COUNT(*) FROM recipes").fetchone()[0]
    fts_count = conn.execute("SELECT COUNT(*) FROM recipes_fts").fetchone()[0]
    conn.close()

    print()
    print(f"Total recipes found: {total_recipes}")
    print(f"Database rows:       {count}")
    print(f"FTS index rows:      {fts_count}")
    print(f"Database saved to:   {db_path}")

    # Export JSON for static site
    print()
    os.makedirs(os.path.dirname(json_path), exist_ok=True)
    export_json(db_path, json_path)


if __name__ == "__main__":
    main()
