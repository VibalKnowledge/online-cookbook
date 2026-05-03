#!/usr/bin/env python3
"""
parse_recipes.py — Parse all .docx recipe files from Joanne's Cookbook
into a SQLite database with full-text search support.

Usage:
    python3 app/parse_recipes.py
"""

import os
import re
import sqlite3
import sys

import docx
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
COOKBOOK_DIR = os.path.join(os.path.dirname(__file__), "..", "Joanne_s Cookbook")
DB_PATH = os.path.join(os.path.dirname(__file__), "..", "recipes.db")

SKIP_FILES = {
    "open file for future use.docx",
    "open space for future additions.docx",
}

# Threshold in EMU – 16 pt = 203200 EMU.  Section headers are ≥ 16 pt bold
# but *not* underlined.
SECTION_HEADER_SIZE_THRESHOLD = 190500  # ~15 pt — safe lower bound

# ---------------------------------------------------------------------------
# Database helpers
# ---------------------------------------------------------------------------

def init_db(db_path: str) -> sqlite3.Connection:
    """Create (or recreate) the database and return a connection."""
    if os.path.exists(db_path):
        os.remove(db_path)
    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.executescript(
        """
        CREATE TABLE recipes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            category TEXT,
            subcategory TEXT,
            ingredients TEXT,
            instructions TEXT,
            notes TEXT,
            serves TEXT,
            source_file TEXT
        );

        CREATE VIRTUAL TABLE recipes_fts USING fts5(
            title, category, subcategory, ingredients, instructions, notes,
            content='recipes', content_rowid='id'
        );

        -- Triggers to keep FTS in sync
        CREATE TRIGGER recipes_ai AFTER INSERT ON recipes BEGIN
            INSERT INTO recipes_fts(rowid, title, category, subcategory,
                                    ingredients, instructions, notes)
            VALUES (new.id, new.title, new.category, new.subcategory,
                    new.ingredients, new.instructions, new.notes);
        END;

        CREATE TRIGGER recipes_ad AFTER DELETE ON recipes BEGIN
            INSERT INTO recipes_fts(recipes_fts, rowid, title, category,
                                    subcategory, ingredients, instructions, notes)
            VALUES ('delete', old.id, old.title, old.category, old.subcategory,
                    old.ingredients, old.instructions, old.notes);
        END;

        CREATE TRIGGER recipes_au AFTER UPDATE ON recipes BEGIN
            INSERT INTO recipes_fts(recipes_fts, rowid, title, category,
                                    subcategory, ingredients, instructions, notes)
            VALUES ('delete', old.id, old.title, old.category, old.subcategory,
                    old.ingredients, old.instructions, old.notes);
            INSERT INTO recipes_fts(rowid, title, category, subcategory,
                                    ingredients, instructions, notes)
            VALUES (new.id, new.title, new.category, new.subcategory,
                    new.ingredients, new.instructions, new.notes);
        END;
        """
    )
    return conn


def insert_recipe(conn, recipe: dict):
    conn.execute(
        """INSERT INTO recipes
           (title, category, subcategory, ingredients, instructions, notes, serves, source_file)
           VALUES (:title, :category, :subcategory, :ingredients, :instructions, :notes, :serves, :source_file)""",
        recipe,
    )


# ---------------------------------------------------------------------------
# Document element iteration (handles interleaved paragraphs + tables)
# ---------------------------------------------------------------------------

def iter_body_elements(doc):
    """Yield (type, element) tuples in document-body order.

    type is 'p' for paragraph or 'tbl' for table.
    For paragraphs, element is a docx.text.paragraph.Paragraph.
    For tables, element is a docx.table.Table.
    """
    body = doc.element.body
    # Build lookup maps so we can wrap raw XML elements.
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p" and child in para_map:
            yield ("p", para_map[child])
        elif tag == "tbl" and child in table_map:
            yield ("tbl", table_map[child])


# ---------------------------------------------------------------------------
# Paragraph classification helpers
# ---------------------------------------------------------------------------

def is_recipe_title(para) -> bool:
    """A recipe title has ALL runs bold AND underlined (checking run-level formatting).

    Excludes:
      - empty / whitespace-only paragraphs
      - section headers (bold, large font, but NOT underlined)
    """
    text = para.text.strip().rstrip("_").strip()
    if not text:
        return False

    runs_with_text = [r for r in para.runs if r.text.strip().rstrip("_").strip()]
    if not runs_with_text:
        return False

    for run in runs_with_text:
        # Must be bold
        if not run.bold:
            return False
        # Must be underlined (True or a WdUnderline enum value)
        if not run.underline:
            return False
        # Reject section headers: bold + underline but very large font
        if run.font.size and run.font.size >= SECTION_HEADER_SIZE_THRESHOLD:
            return False

    return True


def is_section_header(para) -> bool:
    """Section headers are bold, large font (≥16pt), but NOT underlined."""
    text = para.text.strip()
    if not text:
        return False
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if not runs_with_text:
        return False
    for run in runs_with_text:
        if not run.bold:
            return False
        if run.font.size and run.font.size >= SECTION_HEADER_SIZE_THRESHOLD:
            return True
    return False


def clean_title(text: str) -> str:
    """Strip trailing underscores and whitespace from a recipe title."""
    return text.strip().rstrip("_").strip()


def is_ingredient_line(text: str) -> bool:
    """Lines starting with -, •, or tab-indented dashes."""
    stripped = text.lstrip()
    if stripped.startswith("-") or stripped.startswith("•"):
        return True
    # Also catch lines starting with dash after whitespace/tab
    if stripped.startswith("\t-") or stripped.startswith("\t•"):
        return True
    return False


_INGREDIENT_LABEL_RE = re.compile(
    r"^Ingredients\s*[;:]", re.IGNORECASE
)

_SERVES_RE = re.compile(
    r"(?:serves|makes|yield[s]?)\s+(.+)", re.IGNORECASE
)

_NOTE_RE = re.compile(
    r"^(?:NOTES?|SUBSTITUT(?:ION|ITION)S?)\s*[;:.]", re.IGNORECASE
)


def extract_serves(text: str) -> str | None:
    """Pull serving info from an Ingredients: line."""
    m = _SERVES_RE.search(text)
    if m:
        return m.group(1).strip().rstrip(".")
    return None


# ---------------------------------------------------------------------------
# Table text extraction
# ---------------------------------------------------------------------------

def table_to_lines(table) -> list[str]:
    """Convert a Word table into pseudo-ingredient lines.

    Each row becomes one line, cells joined by spaces.
    """
    lines = []
    for row in table.rows:
        parts = [cell.text.strip() for cell in row.cells]
        # Deduplicate adjacent identical cells (merged cells repeat text)
        deduped = []
        for p in parts:
            if not deduped or p != deduped[-1]:
                deduped.append(p)
        line = " ".join(deduped).strip()
        if line:
            lines.append("- " + line)
    return lines


# ---------------------------------------------------------------------------
# Core parser
# ---------------------------------------------------------------------------

def parse_docx(filepath: str, category: str, subcategory: str) -> list[dict]:
    """Parse a single .docx file and return a list of recipe dicts."""
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"  WARNING: Could not open {filepath}: {e}")
        return []

    # Collect all content in body order (paragraphs + table text)
    elements: list[tuple[str, object]] = []  # ('p', para) or ('text', str)
    for etype, elem in iter_body_elements(doc):
        if etype == "p":
            elements.append(("p", elem))
        elif etype == "tbl":
            for line in table_to_lines(elem):
                elements.append(("text", line))

    # -------------------------------------------------------------------
    # Pass 1: Identify title boundaries
    # -------------------------------------------------------------------
    title_indices: list[int] = []
    for i, (etype, elem) in enumerate(elements):
        if etype == "p" and is_recipe_title(elem):
            title_indices.append(i)

    if not title_indices:
        return []

    # -------------------------------------------------------------------
    # Pass 2: Extract recipe blocks between titles
    # -------------------------------------------------------------------
    recipes = []
    for idx, ti in enumerate(title_indices):
        etype, elem = elements[ti]
        title = clean_title(elem.text)
        if not title:
            continue

        # Collect content paragraphs from after title to next title (or end)
        end = title_indices[idx + 1] if idx + 1 < len(title_indices) else len(elements)
        content_lines: list[str] = []
        for j in range(ti + 1, end):
            ctype, celem = elements[j]
            if ctype == "p":
                content_lines.append(celem.text)
            else:  # "text" from table
                content_lines.append(celem)

        recipe = _parse_recipe_content(title, content_lines, category, subcategory, filepath)
        recipes.append(recipe)

    return recipes


def _parse_recipe_content(
    title: str,
    lines: list[str],
    category: str,
    subcategory: str,
    filepath: str,
) -> dict:
    """Classify content lines into ingredients, instructions, and notes."""
    ingredients: list[str] = []
    instructions: list[str] = []
    notes: list[str] = []
    serves: str | None = None

    # State machine: 'pre' (before ingredients), 'ing' (ingredients section),
    # 'inst' (instructions), 'notes'
    state = "pre"

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            # Blank lines can signal end of ingredients section
            if state == "ing":
                state = "inst"
            continue

        # Check for Ingredients: label
        if _INGREDIENT_LABEL_RE.match(line):
            s = extract_serves(line)
            if s:
                serves = s
            state = "ing"
            continue

        # Check for notes
        if _NOTE_RE.match(line):
            notes.append(line)
            state = "notes"
            continue

        # Also catch SUBSTITUTIONS / SUBSTITITIONS (common typo)
        if line.upper().startswith("SUBSTITUT") or line.upper().startswith("SUBSTITUITI"):
            notes.append(line)
            state = "notes"
            continue

        if state == "notes":
            # Continue collecting notes until we hit something that looks
            # like a new section
            if is_ingredient_line(line) or _INGREDIENT_LABEL_RE.match(line):
                state = "ing"
            else:
                notes.append(line)
                continue

        if state in ("pre", "ing"):
            if is_ingredient_line(line):
                state = "ing"
                ingredients.append(line)
                continue
            elif state == "ing":
                # Non-ingredient line after ingredients -> instructions
                state = "inst"
                instructions.append(line)
                continue
            else:
                # Pre-ingredient content (intro text) goes to instructions
                instructions.append(line)
                continue

        if state == "inst":
            if is_ingredient_line(line):
                # Some recipes have a second ingredient group (e.g. topping)
                ingredients.append(line)
                continue
            instructions.append(line)

    return {
        "title": title,
        "category": category,
        "subcategory": subcategory,
        "ingredients": "\n".join(ingredients) if ingredients else None,
        "instructions": "\n".join(instructions) if instructions else None,
        "notes": "\n".join(notes) if notes else None,
        "serves": serves,
        "source_file": os.path.relpath(filepath, os.path.join(os.path.dirname(__file__), "..")),
    }


# ---------------------------------------------------------------------------
# File discovery
# ---------------------------------------------------------------------------

def discover_docx_files(base_dir: str) -> list[tuple[str, str, str]]:
    """Return list of (filepath, category, subcategory) tuples."""
    results = []
    base_dir = os.path.abspath(base_dir)

    for root, _dirs, files in os.walk(base_dir):
        for fname in sorted(files):
            if not fname.lower().endswith(".docx"):
                continue
            if fname.lower() in {s.lower() for s in SKIP_FILES}:
                continue

            filepath = os.path.join(root, fname)
            rel = os.path.relpath(root, base_dir)

            if rel == ".":
                # Files at the top level of the cookbook directory
                category = ""
            else:
                # Use the top-level folder as category (handles nested like Poultry/Poultry)
                parts = rel.split(os.sep)
                category = parts[0].strip()

            subcategory = os.path.splitext(fname)[0].strip()
            results.append((filepath, category, subcategory))

    return results


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    base_dir = os.path.abspath(COOKBOOK_DIR)
    db_path = os.path.abspath(DB_PATH)

    print(f"Cookbook directory: {base_dir}")
    print(f"Database path:     {db_path}")
    print()

    files = discover_docx_files(base_dir)
    print(f"Found {len(files)} .docx files to parse.")
    print()

    conn = init_db(db_path)
    total_recipes = 0

    for i, (filepath, category, subcategory) in enumerate(files, 1):
        short = os.path.relpath(filepath, base_dir)
        recipes = parse_docx(filepath, category, subcategory)
        for r in recipes:
            insert_recipe(conn, r)
        total_recipes += len(recipes)
        print(f"  [{i:3d}/{len(files)}] {short:<65s}  → {len(recipes):3d} recipes")

    conn.commit()

    # Verify
    count = conn.execute("SELECT COUNT(*) FROM recipes").fetchone()[0]
    fts_count = conn.execute("SELECT COUNT(*) FROM recipes_fts").fetchone()[0]
    conn.close()

    print()
    print(f"Total recipes found: {total_recipes}")
    print(f"Database rows:       {count}")
    print(f"FTS index rows:      {fts_count}")
    print(f"Database saved to:   {db_path}")


if __name__ == "__main__":
    main()
